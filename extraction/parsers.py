from __future__ import annotations

import csv
from pathlib import Path
from typing import Any

from extraction.schemas import ParsedDocument

# Threshold for "image-likely" PDF detection. A PDF whose average
# extractable-text density falls below this routes to the native
# PDF document block path (Bedrock vision) instead of the text path.
# Tuned 2026-05-03 against Croesus printscreen exports: a typical
# scanned KYC carries 0-40 chars/page (header watermark only); a
# typical text PDF carries 800-3000 chars/page.
IMAGE_PDF_AVG_CHARS_THRESHOLD = 50

# Companion threshold: ratio of pages with ANY extractable text to
# total pages. Below this we treat the PDF as image-heavy regardless
# of the average. Catches the Croesus 10-page printscreen with one
# text cover page (avg might exceed 50, ratio is 0.1).
IMAGE_PDF_TEXT_PAGE_RATIO_THRESHOLD = 0.5


class ParserDependencyError(RuntimeError):
    pass


def is_likely_image_pdf(parsed: ParsedDocument) -> bool:
    """True if ``parsed`` came from a scanned / image-only / sparse PDF.

    Used by ``extraction.pipeline`` to dispatch real-derived extraction
    to the native PDF document block path (Bedrock vision) instead of
    the text-only path. Returns False for non-PDF inputs and for
    text-rich PDFs.

    Signals:
      1. ``method == "ocr_required"`` — pymupdf returned zero text on
         every page (pure scan).
      2. ``text_page_count / page_count < IMAGE_PDF_TEXT_PAGE_RATIO_THRESHOLD``
         — fewer than half the pages have any extractable text
         (catches Croesus printscreens with metadata-only cover page).
      3. average chars/page < ``IMAGE_PDF_AVG_CHARS_THRESHOLD`` —
         total extracted text is sparse relative to page count
         (catches very low-density text like single-line headers).

    The thresholds are deliberately conservative: false-positives
    route text-rich PDFs to the more expensive vision path. False-
    negatives leave Croesus printscreens stuck on the text path
    where they previously returned 0 facts. Re-tune via canary if
    sweep data shows either edge dominating.
    """
    if parsed.method == "ocr_required":
        return True
    if parsed.method != "pdf_native":
        return False
    metadata = parsed.metadata or {}
    page_count = int(metadata.get("page_count") or 0)
    if page_count <= 0:
        return False
    text_page_count = int(metadata.get("text_page_count") or 0)
    if text_page_count / page_count < IMAGE_PDF_TEXT_PAGE_RATIO_THRESHOLD:
        return True
    total_chars = len((parsed.text or "").strip())
    avg_chars_per_page = total_chars / page_count
    return avg_chars_per_page < IMAGE_PDF_AVG_CHARS_THRESHOLD


def parse_document_path(path: Path) -> ParsedDocument:
    extension = path.suffix.lower()
    if extension == ".pdf":
        return _parse_pdf(path)
    if extension == ".docx":
        return _parse_docx(path)
    if extension == ".xlsx":
        return _parse_xlsx(path)
    if extension == ".csv":
        return _parse_csv(path)
    if extension in {".txt", ".md"}:
        return ParsedDocument(path.read_text(errors="ignore"), "plain", {"extension": extension})
    if extension in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        return ParsedDocument("", "ocr_required", {"reason": "image_file", "extension": extension})
    return ParsedDocument("", "unsupported", {"extension": extension})


def _parse_pdf(path: Path) -> ParsedDocument:
    try:
        import fitz
    except ImportError as exc:
        raise ParserDependencyError("pymupdf is required for PDF parsing.") from exc

    doc = fitz.open(path)
    pages: list[str] = []
    fragments: list[dict[str, Any]] = []
    for index, page in enumerate(doc, start=1):
        page_text = page.get_text("text")
        if page_text:
            pages.append(f"[page {index}]\n{page_text}")
            fragments.append({"kind": "page_text", "page": index, "char_count": len(page_text)})
    text = "\n\n".join(pages)
    metadata = {"page_count": doc.page_count, "text_page_count": len(pages)}
    if not text.strip():
        return ParsedDocument("", "ocr_required", metadata, fragments)
    return ParsedDocument(text, "pdf_native", metadata, fragments)


def _parse_docx(path: Path) -> ParsedDocument:
    try:
        import docx
    except ImportError as exc:
        raise ParserDependencyError("python-docx is required for DOCX parsing.") from exc

    document = docx.Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    tables: list[str] = []
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                tables.append(f"[table {table_index} row {row_index}] " + " | ".join(values))
    return ParsedDocument(
        "\n".join([*paragraphs, *tables]),
        "docx",
        {"paragraph_count": len(paragraphs), "table_rows": len(tables)},
    )


def _parse_xlsx(path: Path) -> ParsedDocument:
    try:
        import openpyxl
    except ImportError as exc:
        raise ParserDependencyError("openpyxl is required for XLSX parsing.") from exc

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    chunks: list[str] = []
    fragments: list[dict[str, Any]] = []
    sheet_names = [sheet.title for sheet in workbook.worksheets]
    for sheet in workbook.worksheets:
        chunks.append(f"[sheet {sheet.title}]")
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = [str(value) for value in row if value is not None and value != ""]
            if values:
                chunks.append(f"[{sheet.title}!{row_index}] " + " | ".join(values))
                fragments.append({"kind": "sheet_row", "sheet": sheet.title, "row": row_index})
    return ParsedDocument(
        "\n".join(chunks),
        "xlsx",
        {"sheet_count": len(workbook.worksheets), "sheet_names": sheet_names},
        fragments,
    )


def _parse_csv(path: Path) -> ParsedDocument:
    rows: list[str] = []
    with path.open(errors="ignore", newline="") as handle:
        reader = csv.reader(handle)
        for index, row in enumerate(reader, start=1):
            values = [value.strip() for value in row if value.strip()]
            if values:
                rows.append(f"[row {index}] " + " | ".join(values))
    return ParsedDocument("\n".join(rows), "csv", {"row_count": len(rows)})
