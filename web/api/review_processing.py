from __future__ import annotations

import base64
import json
import mimetypes
import os
import re
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Literal

from django.conf import settings
from django.core.exceptions import ImproperlyConfigured
from django.db import transaction
from django.utils import timezone
from pydantic import BaseModel, Field, ValidationError

from web.api import models
from web.api.review_redaction import (
    pii_detection_summary,
    redact_evidence_quote,
    sanitize_fact_value,
)
from web.api.review_security import secure_data_root
from web.api.review_state import create_state_version, reviewed_state_from_workspace
from web.audit.writer import record_event

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".xlsx",
    ".csv",
    ".txt",
    ".md",
    ".png",
    ".jpg",
    ".jpeg",
    ".tif",
    ".tiff",
}

DOCUMENT_TYPE_BY_HINT = {
    "kyc": "kyc",
    "know your client": "kyc",
    "statement": "statement",
    "holdings": "statement",
    "plan": "planning",
    "projection": "planning",
    "croesus": "crm_export",
    "crm": "crm_export",
    "intake": "intake",
    "meeting": "meeting_note",
    "note": "meeting_note",
}


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    method: str
    metadata: dict[str, Any]


class BedrockFact(BaseModel):
    field: str = Field(min_length=1)
    value: Any
    confidence: Literal["high", "medium", "low"] = "medium"
    derivation_method: Literal["extracted", "inferred", "defaulted"] = "extracted"
    source_location: str = ""
    source_page: int | None = None
    evidence_quote: str = ""


class BedrockFactsPayload(BaseModel):
    facts: list[BedrockFact] = Field(default_factory=list)


def record_worker_heartbeat(
    *,
    name: str | None = None,
    current_job: models.ProcessingJob | None = None,
    metadata: dict[str, Any] | None = None,
) -> models.WorkerHeartbeat:
    heartbeat, _ = models.WorkerHeartbeat.objects.update_or_create(
        name=name or settings.MP20_WORKER_NAME,
        defaults={
            "last_seen_at": timezone.now(),
            "current_job": current_job,
            "metadata": metadata or {},
        },
    )
    return heartbeat


def claim_next_job() -> models.ProcessingJob | None:
    with transaction.atomic():
        job = (
            models.ProcessingJob.objects.select_for_update(skip_locked=True)
            .filter(status=models.ProcessingJob.Status.QUEUED)
            .order_by("created_at")
            .first()
        )
        if job is None:
            return None
        job.status = models.ProcessingJob.Status.PROCESSING
        job.attempts += 1
        job.locked_at = timezone.now()
        job.started_at = job.started_at or timezone.now()
        job.metadata = {**job.metadata, "stage": "claimed"}
        job.save(
            update_fields=[
                "status",
                "attempts",
                "locked_at",
                "started_at",
                "metadata",
                "updated_at",
            ]
        )
        return job


def process_job(job: models.ProcessingJob) -> None:
    record_worker_heartbeat(current_job=job, metadata={"stage": "processing"})
    try:
        job.metadata = {**job.metadata, "stage": job.job_type}
        job.save(update_fields=["metadata", "updated_at"])
        if job.job_type == models.ProcessingJob.JobType.PROCESS_DOCUMENT:
            if job.document is None:
                raise ValueError("process_document job requires a document")
            process_document(job.document)
        elif job.job_type == models.ProcessingJob.JobType.RECONCILE_WORKSPACE:
            reconcile_workspace(job.workspace)
        else:
            raise ValueError(f"Unsupported job type: {job.job_type}")
    except Exception as exc:
        _fail_or_retry(job, exc)
        return

    job.status = models.ProcessingJob.Status.COMPLETED
    job.completed_at = timezone.now()
    job.last_error = ""
    job.metadata = {**job.metadata, "stage": "completed"}
    job.save(update_fields=["status", "completed_at", "last_error", "metadata", "updated_at"])
    record_worker_heartbeat(current_job=None, metadata={"stage": "idle"})


def process_document(document: models.ReviewDocument) -> None:
    secure_data_root()
    extension = Path(document.original_filename).suffix.lower()
    document.extension = extension.lstrip(".")
    document.document_type = classify_document(document.original_filename, extension)
    document.status = models.ReviewDocument.Status.CLASSIFIED
    document.save(update_fields=["extension", "document_type", "status", "updated_at"])

    if extension not in SUPPORTED_EXTENSIONS:
        document.status = models.ReviewDocument.Status.UNSUPPORTED
        document.failure_reason = f"Unsupported file type: {extension or '[none]'}"
        document.save(update_fields=["status", "failure_reason", "updated_at"])
        record_event(
            action="review_document_unsupported",
            entity_type="review_document",
            entity_id=str(document.id),
            metadata={"extension": extension},
        )
        return

    parsed = parse_document(document)
    document.status = (
        models.ReviewDocument.Status.OCR_REQUIRED
        if parsed.method == "ocr_required"
        else models.ReviewDocument.Status.TEXT_EXTRACTED
    )
    document.processing_metadata = {
        **document.processing_metadata,
        "parse_method": parsed.method,
        "parse_metadata": parsed.metadata,
        "pii_detection_summary": pii_detection_summary(parsed.text),
    }
    document.save(update_fields=["status", "processing_metadata", "updated_at"])

    if (
        not parsed.text.strip()
        and document.workspace.data_origin == models.ReviewWorkspace.DataOrigin.REAL_DERIVED
    ):
        ensure_bedrock_configured()

    facts = extract_facts(document, parsed.text)
    models.ExtractedFact.objects.filter(document=document).delete()
    for fact in facts:
        models.ExtractedFact.objects.create(
            workspace=document.workspace,
            document=document,
            field=fact["field"],
            value=sanitize_fact_value(fact["field"], fact["value"]),
            asserted_at=fact.get("asserted_at") or None,
            confidence=fact.get("confidence", "medium"),
            derivation_method=fact.get("derivation_method", "extracted"),
            source_page=fact.get("source_page"),
            source_location=fact.get("source_location", ""),
            evidence_quote=redact_evidence_quote(fact.get("evidence_quote", "")),
            extraction_run_id=fact["extraction_run_id"],
        )

    document.status = models.ReviewDocument.Status.FACTS_EXTRACTED
    document.save(update_fields=["status", "updated_at"])
    enqueue_reconcile(document.workspace)
    record_event(
        action="review_document_processed",
        entity_type="review_document",
        entity_id=str(document.id),
        metadata={"document_type": document.document_type, "fact_count": len(facts)},
    )


def reconcile_workspace(workspace: models.ReviewWorkspace) -> dict[str, Any]:
    state = reviewed_state_from_workspace(workspace)
    workspace.reviewed_state = state
    workspace.readiness = state["readiness"]
    workspace.status = (
        models.ReviewWorkspace.Status.ENGINE_READY
        if state["readiness"]["engine_ready"]
        else models.ReviewWorkspace.Status.REVIEW_READY
    )
    workspace.match_candidates = []
    workspace.save(
        update_fields=["reviewed_state", "readiness", "status", "match_candidates", "updated_at"]
    )
    create_state_version(workspace, user=None, state=state)
    workspace.documents.filter(status=models.ReviewDocument.Status.FACTS_EXTRACTED).update(
        status=models.ReviewDocument.Status.RECONCILED
    )
    record_event(
        action="review_workspace_reconciled",
        entity_type="review_workspace",
        entity_id=workspace.external_id,
        metadata={
            "workspace_id": workspace.external_id,
            "engine_ready": state["readiness"]["engine_ready"],
            "missing_count": len(state["readiness"].get("missing", [])),
        },
    )
    return state


def enqueue_reconcile(workspace: models.ReviewWorkspace) -> models.ProcessingJob:
    pending = models.ProcessingJob.objects.filter(
        workspace=workspace,
        job_type=models.ProcessingJob.JobType.RECONCILE_WORKSPACE,
        status__in=[models.ProcessingJob.Status.QUEUED, models.ProcessingJob.Status.PROCESSING],
    ).first()
    if pending:
        return pending
    return models.ProcessingJob.objects.create(
        workspace=workspace,
        job_type=models.ProcessingJob.JobType.RECONCILE_WORKSPACE,
        metadata={"stage": "queued_reconcile"},
    )


def classify_document(filename: str, extension: str) -> str:
    lowered = filename.lower().replace("_", " ")
    for hint, document_type in DOCUMENT_TYPE_BY_HINT.items():
        if hint in lowered:
            return document_type
    if extension == ".xlsx":
        return "spreadsheet"
    if extension == ".csv":
        return "spreadsheet"
    if extension in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        return "image"
    return "unknown"


def parse_document(document: models.ReviewDocument) -> ParsedDocument:
    path = secure_data_root() / document.storage_path
    extension = path.suffix.lower()
    if extension == ".pdf":
        return _parse_pdf(path)
    if extension == ".docx":
        return _parse_docx(path)
    if extension == ".xlsx":
        return _parse_xlsx(path)
    if extension == ".csv":
        return ParsedDocument(path.read_text(errors="ignore"), "csv", {})
    if extension in {".txt", ".md"}:
        return ParsedDocument(path.read_text(errors="ignore"), "plain", {})
    if extension in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        return ParsedDocument("", "ocr_required", {"reason": "image_file"})
    return ParsedDocument("", "unsupported", {"extension": extension})


def extract_facts(document: models.ReviewDocument, text: str) -> list[dict[str, Any]]:
    extraction_run_id = f"{document.document_type}:{int(time.time())}"
    if document.workspace.data_origin == models.ReviewWorkspace.DataOrigin.REAL_DERIVED:
        ensure_bedrock_configured()
        if text.strip():
            return _bedrock_extract_facts(document, text, extraction_run_id)
        return _bedrock_extract_visual_facts(document, extraction_run_id)
    return _heuristic_facts(document, text, extraction_run_id)


def ensure_bedrock_configured() -> None:
    missing = [
        name
        for name in ("AWS_ACCESS_KEY_ID", "AWS_SECRET_ACCESS_KEY", "BEDROCK_MODEL")
        if not os.getenv(name)
    ]
    if missing:
        raise ImproperlyConfigured(
            "Real-derived extraction requires Bedrock configuration: " + ", ".join(missing)
        )


def _bedrock_extract_facts(
    document: models.ReviewDocument, text: str, extraction_run_id: str
) -> list[dict[str, Any]]:
    client = _bedrock_client()
    limit = getattr(settings, "MP20_TEXT_EXTRACTION_MAX_CHARS", 24000)
    prompt = _fact_extraction_prompt(document=document, text=text[:limit])
    response = client.messages.create(
        model=os.getenv("BEDROCK_MODEL", settings.BEDROCK_MODEL),
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )
    return _facts_from_bedrock_response(response, extraction_run_id)


def _bedrock_extract_visual_facts(
    document: models.ReviewDocument, extraction_run_id: str
) -> list[dict[str, Any]]:
    client = _bedrock_client()
    image_blocks, overflow = _visual_content_blocks(document)
    if not image_blocks:
        raise ValueError(
            "Document requires OCR, but no supported visual payload could be prepared."
        )
    if overflow:
        document.processing_metadata = {
            **document.processing_metadata,
            "ocr_overflow": overflow,
        }
        document.save(update_fields=["processing_metadata", "updated_at"])
    prompt = _fact_extraction_prompt(
        document=document,
        text="Use the attached page or image content as the source.",
    )
    response = client.messages.create(
        model=os.getenv("BEDROCK_MODEL", settings.BEDROCK_MODEL),
        max_tokens=4096,
        messages=[{"role": "user", "content": [*image_blocks, {"type": "text", "text": prompt}]}],
    )
    return _facts_from_bedrock_response(response, extraction_run_id)


def _bedrock_client():
    try:
        import anthropic
    except ImportError as exc:
        raise ImproperlyConfigured(
            "anthropic[bedrock] is required for Bedrock extraction."
        ) from exc

    return anthropic.AnthropicBedrock(
        aws_access_key=os.getenv("AWS_ACCESS_KEY_ID"),
        aws_secret_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
        aws_region=os.getenv("AWS_REGION", settings.AWS_REGION),
    )


def _fact_extraction_prompt(*, document: models.ReviewDocument, text: str) -> str:
    return (
        "Extract MP2.0 review facts from this client document. Return JSON only: "
        '{"facts":[{"field":"...","value":...,"confidence":"high|medium|low",'
        '"derivation_method":"extracted|inferred|defaulted","source_location":"...",'
        '"source_page":null,"evidence_quote":"short quote"}]}. '
        "Use fields like household.display_name, people, accounts, goals, "
        "goal_account_links, risk.household_score. Do not invent missing values. "
        "For account numbers, SIN, SSN, tax IDs, and similar identifiers, include only "
        "the field name and raw value in JSON; the application will hash and redact them.\n\n"
        f"Document type: {document.document_type}\nText:\n{text}"
    )


def _facts_from_bedrock_response(response, extraction_run_id: str) -> list[dict[str, Any]]:  # noqa: ANN001
    content = "".join(block.text for block in response.content if hasattr(block, "text"))
    payload = _json_payload_from_model_text(content)
    try:
        parsed = BedrockFactsPayload.model_validate(payload)
    except ValidationError as exc:
        raise ValueError("Bedrock extraction JSON did not match the expected fact schema.") from exc
    facts = []
    for fact in parsed.facts:
        item = fact.model_dump(mode="json")
        item["extraction_run_id"] = extraction_run_id
        facts.append(item)
    return facts


def _json_payload_from_model_text(content: str) -> dict[str, Any]:
    normalized = content.strip()
    normalized = re.sub(r"^```(?:json)?\s*", "", normalized)
    normalized = re.sub(r"\s*```$", "", normalized)
    try:
        return json.loads(normalized)
    except json.JSONDecodeError as exc:
        repaired = re.sub(r",(\s*[}\]])", r"\1", normalized)
        try:
            return json.loads(repaired)
        except json.JSONDecodeError:
            pass
        start = normalized.find("{")
        end = normalized.rfind("}")
        if start == -1 or end == -1 or end <= start:
            raise ValueError("Bedrock extraction did not return valid JSON.") from exc
        try:
            return json.loads(re.sub(r",(\s*[}\]])", r"\1", normalized[start : end + 1]))
        except json.JSONDecodeError as nested_exc:
            raise ValueError("Bedrock extraction did not return valid JSON.") from nested_exc


def _visual_content_blocks(
    document: models.ReviewDocument,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    path = secure_data_root() / document.storage_path
    extension = path.suffix.lower()
    if extension == ".pdf":
        return _pdf_page_image_blocks(path)

    media_type = mimetypes.guess_type(path.name)[0] or ""
    if media_type not in {"image/png", "image/jpeg", "image/gif", "image/webp"}:
        raise ValueError(
            f"Bedrock vision OCR does not support {extension or 'this file type'} yet."
        )
    return [_image_block(path.read_bytes(), media_type)], {}


def _pdf_page_image_blocks(path: Path) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    try:
        import fitz
    except ImportError as exc:
        raise ImproperlyConfigured("pymupdf is required for PDF OCR preparation.") from exc

    document = fitz.open(path)
    blocks: list[dict[str, Any]] = []
    max_pages = getattr(settings, "MP20_OCR_MAX_PAGES", 12)
    for page in list(document)[:max_pages]:
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        blocks.append(_image_block(pixmap.tobytes("png"), "image/png"))
    overflow = {}
    if document.page_count > max_pages:
        overflow = {
            "total_pages": document.page_count,
            "processed_pages": max_pages,
            "overflow_pages": document.page_count - max_pages,
            "status": "needs_review",
        }
    return blocks, overflow


def _image_block(content: bytes, media_type: str) -> dict[str, Any]:
    return {
        "type": "image",
        "source": {
            "type": "base64",
            "media_type": media_type,
            "data": base64.b64encode(content).decode(),
        },
    }


def _heuristic_facts(
    document: models.ReviewDocument, text: str, extraction_run_id: str
) -> list[dict[str, Any]]:
    facts: list[dict[str, Any]] = []
    stripped = " ".join(text.split())
    if stripped:
        facts.append(
            {
                "field": f"document.{document.id}.summary",
                "value": stripped[:500],
                "confidence": "low",
                "derivation_method": "extracted",
                "source_location": "document",
                "evidence_quote": stripped[:240],
                "extraction_run_id": extraction_run_id,
            }
        )
    if document.document_type == "statement":
        account_hint = {
            "id": f"review_account_{document.id}",
            "type": "Non-Registered",
            "current_value": 0,
            "missing_holdings_confirmed": True,
        }
        facts.append(
            {
                "field": "accounts",
                "value": [account_hint],
                "confidence": "low",
                "derivation_method": "inferred",
                "source_location": "filename",
                "evidence_quote": document.original_filename,
                "extraction_run_id": extraction_run_id,
            }
        )
    if document.document_type == "crm_export":
        display_name = Path(document.original_filename).stem.replace("_", " ").replace("-", " ")
        facts.append(
            {
                "field": "household.display_name",
                "value": display_name,
                "confidence": "low",
                "derivation_method": "inferred",
                "source_location": "filename",
                "evidence_quote": document.original_filename,
                "extraction_run_id": extraction_run_id,
            }
        )
    return facts


def _parse_pdf(path: Path) -> ParsedDocument:
    try:
        import fitz
    except ImportError as exc:
        raise ImproperlyConfigured("pymupdf is required for PDF parsing.") from exc

    doc = fitz.open(path)
    pages: list[str] = []
    for index, page in enumerate(doc, start=1):
        page_text = page.get_text("text")
        if page_text:
            pages.append(f"[page {index}]\n{page_text}")
    text = "\n\n".join(pages)
    if not text.strip():
        return ParsedDocument("", "ocr_required", {"page_count": doc.page_count})
    return ParsedDocument(text, "pdf_native", {"page_count": doc.page_count})


def _parse_docx(path: Path) -> ParsedDocument:
    try:
        import docx
    except ImportError as exc:
        raise ImproperlyConfigured("python-docx is required for DOCX parsing.") from exc

    document = docx.Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    tables: list[str] = []
    for table in document.tables:
        for row in table.rows:
            tables.append(" | ".join(cell.text.strip() for cell in row.cells))
    return ParsedDocument("\n".join([*paragraphs, *tables]), "docx", {"table_rows": len(tables)})


def _parse_xlsx(path: Path) -> ParsedDocument:
    try:
        import openpyxl
    except ImportError as exc:
        raise ImproperlyConfigured("openpyxl is required for XLSX parsing.") from exc

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    chunks: list[str] = []
    for sheet in workbook.worksheets:
        chunks.append(f"[sheet {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value) for value in row if value is not None and value != ""]
            if values:
                chunks.append(" | ".join(values))
    return ParsedDocument("\n".join(chunks), "xlsx", {"sheet_count": len(workbook.worksheets)})


def _fail_or_retry(job: models.ProcessingJob, exc: Exception) -> None:
    job.last_error = str(exc)
    job.metadata = {
        **job.metadata,
        "stage": job.metadata.get("stage", job.job_type),
        "failure_code": exc.__class__.__name__,
    }
    if job.attempts < job.max_attempts:
        job.status = models.ProcessingJob.Status.QUEUED
    else:
        job.status = models.ProcessingJob.Status.FAILED
        job.completed_at = timezone.now()
        if job.document:
            job.document.status = models.ReviewDocument.Status.FAILED
            job.document.failure_reason = str(exc)
            job.document.processing_metadata = {
                **job.document.processing_metadata,
                "failure_code": exc.__class__.__name__,
                "failure_stage": job.metadata.get("stage", job.job_type),
            }
            job.document.save(
                update_fields=[
                    "status",
                    "failure_reason",
                    "processing_metadata",
                    "updated_at",
                ]
            )
    job.save(update_fields=["status", "last_error", "completed_at", "metadata", "updated_at"])
    record_event(
        action="review_processing_failed",
        entity_type="processing_job",
        entity_id=str(job.id),
        metadata={
            "workspace_id": job.workspace.external_id,
            "attempts": job.attempts,
            "max_attempts": job.max_attempts,
            "will_retry": job.status == models.ProcessingJob.Status.QUEUED,
            "failure_code": exc.__class__.__name__,
        },
    )
    record_worker_heartbeat(current_job=None, metadata={"stage": "idle_after_failure"})
